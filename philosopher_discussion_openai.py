import os
import json
import requests
import time
import random
from typing import List, Dict
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class RateLimiter:
    """ API 速率限制器，每分钟最多允许 N 个请求 """
    def __init__(self, requests_per_minute):
        self.interval = 60.0 / requests_per_minute
        self.last_request = 0

    def wait(self):
        now = time.time()
        elapsed = now - self.last_request
        if elapsed < self.interval:
            time.sleep(self.interval - elapsed)
        self.last_request = time.time()

class PhilosopherDiscussionOpenAI:
    def __init__(self):
        # API 配置
        os.environ['NO_PROXY'] = '*'
        self.url = "https://api.openai.com/v1/chat/completions"
        self.headers = {
            "Authorization": "Bearer ",
            "Content-Type": "application/json"
        }
        
        # 设置代理
        self.proxies = {
            'http': 'http://127.0.0.1:7890',
            'https': 'http://127.0.0.1:7890'
        }
        
        # API 速率限制：每分钟 5 个请求
        self.rate_limiter = RateLimiter(5)

        # 加载哲学家 prompts
        self.philosophers = {
            "aristotle": self._load_prompt("prompts/aristotle.txt"),
            "confucius": self._load_prompt("prompts/confucius.txt"),
            "kant": self._load_prompt("prompts/kant.txt")
        }

        # 加载裁判 prompt
        self.judge_prompt = self._load_prompt("prompts/judge.txt")

    def _load_prompt(self, file_path: str) -> str:
        """ 读取 prompt 文件内容 """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except FileNotFoundError:
            print(f"错误：未找到 {file_path}")
            return ""

    def get_response(self, prompt: str, question: str, retries=5) -> str:
        """ 发送请求并获取 AI 响应，支持重试机制 """
        data = {
            "model": "gpt-4o-mini",
            "messages": [
                {"role": "system", "content": prompt},
                {"role": "user", "content": question}
            ],
            "temperature": 0.5,
            "max_tokens": 1000
        }

        for attempt in range(1, retries + 1):
            try:
                self.rate_limiter.wait()
                
                response = requests.post(
                    self.url, 
                    headers=self.headers, 
                    json=data,
                    proxies=self.proxies,
                    timeout=(60, 120)
                )
                
                response.raise_for_status()
                result = response.json()

                if 'choices' in result and len(result['choices']) > 0:
                    content = result['choices'][0]['message']['content'].strip()
                    print(f"成功获取回答（前100字）：{content[:100]}...")
                    return content

                raise ValueError("API返回格式错误")

            except Exception as e:
                print(f"请求失败（尝试 {attempt}/{retries}）: {e}")
                if attempt < retries:
                    wait_time = min(attempt * 15, 60)
                    print(f"等待 {wait_time} 秒后重试...")
                    time.sleep(wait_time)
                else:
                    return f"API请求失败: {str(e)}"

        return "所有重试都失败了"

    def generate_article(self, discussion_history: List[Dict]) -> str:
        """ 生成总结文章 """
        # 切换到 DeepSeek API
        deepseek_url = "https://api.siliconflow.cn/v1/chat/completions"
        deepseek_headers = {
            "Authorization": "Bearer ",
            "Content-Type": "application/json"
        }

        article_prompt = """
        请将以下哲学讨论整理成一篇严谨的学术文章，要求：
        1. 采用总-分-总的结构
        2. 突出各个哲学家观点的异同
        3. 分析讨论的演进过程
        4. 总结各方观点的价值
        5.不要用列举的方法，应该用通顺的话语将你的论点串联起来
        6.不要用“讨论”、“讨论记录”、“讨论总结”等字眼，应该用“文章”、“论文”等字眼

        讨论内容如下：
        """

        discussion_text = json.dumps(discussion_history, ensure_ascii=False, indent=2)
        
        data = {
            "model": "Pro/deepseek-ai/DeepSeek-R1",  # 使用 DeepSeek-R1 模型
            "messages": [
                {"role": "system", "content": article_prompt},
                {"role": "user", "content": discussion_text}
            ],
            "temperature": 0.3,
            "max_tokens": 4000
        }

        try:
            self.rate_limiter.wait()
            
            response = requests.post(
                deepseek_url,  # 使用 DeepSeek API
                headers=deepseek_headers,
                json=data,
                timeout=(60, 180)
            )
            
            response.raise_for_status()
            result = response.json()

            if 'choices' in result and len(result['choices']) > 0:
                content = result['choices'][0]['message']['content'].strip()
                print("成功生成总结文章")
                return content

            raise ValueError("API返回格式错误")

        except Exception as e:
            print(f"生成总结文章时出错: {e}")
            return f"生成总结文章失败: {str(e)}"

    def conduct_discussion(self, initial_question: str, rounds: int = 3) -> List[Dict]:
        """ 进行哲学家间的多轮讨论 """
        discussion_history = []
        current_question = initial_question
        
        # 从问题中提取主题关键词（取前10个字符）
        topic = initial_question[:10].replace("?", "").replace("？", "").strip()
        
        # 创建讨论记录文件夹
        os.makedirs('discussion_records', exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        discussion_file = f'discussion_records/openai_{topic}_{timestamp}.json'

        def save_discussion():
            """ 保存讨论记录到JSON文件 """
            with open(discussion_file, 'w', encoding='utf-8') as f:
                json.dump({
                    'initial_question': initial_question,
                    'timestamp': timestamp,
                    'rounds': discussion_history
                }, f, ensure_ascii=False, indent=2)

        for round_num in range(rounds):
            print(f"\n第 {round_num + 1} 轮讨论 - 当前问题: {current_question}")
            
            round_responses = {}
            round_data = {
                "round": round_num + 1,
                "question": current_question,
                "responses": {},
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }

            # 获取每个哲学家的回答
            for name, prompt in self.philosophers.items():
                print(f"获取 {name} 的回答...")
                response = self.get_response(prompt, current_question)
                
                if not response.startswith(("API请求失败", "API请求超时")):
                    round_responses[name] = response
                    round_data["responses"][name] = {
                        "content": response,
                        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    }
                    print(f"成功获取 {name} 的回答")
                else:
                    print(f"警告: {name} 的回答获取失败")
                    round_data["responses"][name] = {
                        "error": response,
                        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    }

            if round_responses:
                discussion_history.append(round_data)
                save_discussion()  # 每轮结束后保存

                if round_num < rounds - 1:
                    print("\n裁判正在生成下一轮问题...")
                    judge_input = f"问题：{current_question}\n\n"
                    for name, response in round_responses.items():
                        judge_input += f"{name} 的回答：{response}\n\n"
                    
                    # 为裁判提问设置特殊参数
                    judge_data = {
                        "model": "gpt-4o-mini",
                        "messages": [
                            {"role": "system", "content": self.judge_prompt},
                            {"role": "user", "content": judge_input}
                        ],
                        "temperature": 0.7,
                        "max_tokens": 50,  # 限制新问题的长度
                        "top_p": 0.95
                    }
                    
                    max_retries = 3
                    for retry in range(max_retries):
                        try:
                            self.rate_limiter.wait()
                            response = requests.post(
                                self.url,
                                headers=self.headers,
                                json=judge_data,
                                proxies=self.proxies,
                                timeout=(30, 60)
                            )
                            
                            response.raise_for_status()
                            result = response.json()
                            
                            if 'choices' in result and len(result['choices']) > 0:
                                new_question = result['choices'][0]['message']['content'].strip()
                                # 确保问题不会太长
                                if len(new_question) > 100:
                                    new_question = new_question[:100] + "..."
                                current_question = new_question
                                break
                            
                        except Exception as e:
                            print(f"裁判提问失败，第 {retry + 1} 次重试...")
                            if retry < max_retries - 1:
                                time.sleep(min((retry + 1) * 15, 60))
                            else:
                                print("裁判提问最终失败，使用默认问题继续...")
                                current_question = "请继续探讨上一个问题的深层含义。"
            else:
                print("本轮无有效回答，跳过并继续...")
                round_data["error"] = "本轮无有效回答"
                discussion_history.append(round_data)
                save_discussion()

        return discussion_history

    def generate_word(self, discussion_history: List[Dict], article: str):
        """生成Word格式的讨论记录和总结文章"""
        initial_question = discussion_history[0]['question'] if discussion_history else "未知问题"
        topic = initial_question[:10].replace("?", "").replace("？", "").strip()
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        word_file = f'discussion_records/openai_{topic}_{timestamp}.docx'
        
        # 创建文档
        doc = Document()
        
        # 统一设置默认字体为微软雅黑
        for style in doc.styles:
            if hasattr(style, 'font'):
                style.font.name = '微软雅黑'
        
        # 设置标题样式
        title_style = doc.styles['Title']
        title_style.font.size = Pt(16)
        
        # 设置正文样式
        normal_style = doc.styles['Normal']
        normal_style.font.size = Pt(10.5)
        
        # 设置标题1样式
        heading1_style = doc.styles['Heading 1']
        heading1_style.font.size = Pt(14)
        
        # 设置标题2样式
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.size = Pt(12)
        
        # 添加标题
        doc.add_heading("哲学家讨论记录", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加初始问题
        doc.add_heading("初始问题", 1)
        doc.add_paragraph(initial_question)
        doc.add_paragraph()  # 添加空行
        
        # 添加讨论记录
        for round_data in discussion_history:
            # 添加轮次标题
            doc.add_heading(f"第 {round_data['round']} 轮讨论", 1)
            doc.add_heading("问题", 2)
            doc.add_paragraph(round_data['question'])
            doc.add_paragraph()
            
            # 添加回答
            doc.add_heading("各位哲学家的回答", 2)
            for name, response_data in round_data['responses'].items():
                p = doc.add_paragraph()
                p.add_run(f"{name}的回答：").bold = True
                if 'content' in response_data:
                    p.add_run("\n" + response_data['content'])
                else:
                    p.add_run("\n(回答获取失败: " + response_data.get('error', '未知错误') + ")")
                doc.add_paragraph()
            
            if 'error' in round_data:
                p = doc.add_paragraph()
                p.add_run("注意：" + round_data['error']).font.color.rgb = RGBColor(255, 0, 0)
            
            doc.add_paragraph()  # 添加空行
        
        # 添加总结文章
        doc.add_heading("讨论总结", 1)
        if article and not article.startswith(("API请求失败", "API请求超时")):
            doc.add_paragraph(article)
        else:
            doc.add_paragraph("由于技术原因，未能生成总结文章。以下是讨论要点：")
            summary = self.generate_simple_summary(discussion_history)
            doc.add_paragraph(summary)
        
        # 保存文档
        try:
            doc.save(word_file)
            print(f"Word文件已生成：{word_file}")
        except Exception as e:
            print(f"Word文件生成失败：{e}")
            # 保存为文本文件作为备份
            with open(f'discussion_records/openai_{topic}_{timestamp}.txt', 'w', encoding='utf-8') as f:
                f.write(str(discussion_history))

    def generate_simple_summary(self, discussion_history: List[Dict]) -> str:
        """生成简单的讨论总结"""
        summary = []
        summary.append("讨论要点总结：")
        
        for round_data in discussion_history:
            summary.append(f"\n第{round_data['round']}轮讨论：")
            summary.append(f"问题：{round_data['question']}")
            
            for name, response_data in round_data['responses'].items():
                if 'content' in response_data:
                    # 提取回答的前100个字符作为摘要
                    content = response_data['content'][:100] + "..."
                    summary.append(f"{name}的观点：{content}")
                else:
                    summary.append(f"{name}：(未能获取完整回答)")
        
        return "\n".join(summary)

def load_questions() -> List[str]:
    """从JSON文件加载问题列表"""
    try:
        with open('fixed_simple_question.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
            # 提取所有完整的问题（去除不完整的问题）
            questions = [q['question'].strip('"') for q in data if len(q['question']) < 100 
                        and not q['question'].endswith(('...', '。', '？"', '"'))]
            return questions
    except Exception as e:
        print(f"加载问题失败: {e}")
        return ["什么是真正的幸福？"]  # 返回默认问题

def get_user_input() -> str:
    """获取用户输入的问题或选择随机问题"""
    while True:
        print("\n请选择问题来源：")
        print("1. 从问题库随机选择")
        print("2. 自己输入问题")
        choice = input("请输入选项（1或2）: ").strip()
        
        if choice == "1":
            questions = load_questions()
            question = random.choice(questions)
            print(f"\n随机选择的问题: {question}")
            confirm = input("是否使用这个问题？(y/n): ").lower()
            if confirm == 'y':
                return question
        elif choice == "2":
            question = input("\n请输入您的问题: ").strip()
            if question:
                return question
        else:
            print("无效的选项，请重新选择")

def main():
    discussion = PhilosopherDiscussionOpenAI()
    
    # 获取用户选择的问题
    initial_question = get_user_input()
    print(f"\n开始讨论问题: {initial_question}")

    # 进行讨论
    history = discussion.conduct_discussion(initial_question)

    # 尝试生成总结文章
    article = discussion.generate_article(history)

    # 生成Word文档
    discussion.generate_word(history, article)

    # 保存JSON格式的原始数据
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    topic = initial_question[:10].replace("?", "").replace("？", "").strip()
    
    with open(f'discussion_records/openai_{topic}_{timestamp}.json', 'w', encoding='utf-8') as f:
        json.dump({
            'initial_question': initial_question,
            'timestamp': timestamp,
            'rounds': history,
            'summary_article': article
        }, f, ensure_ascii=False, indent=2)

if __name__ == "__main__":
    main() 